import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def create_element(name):
    return OxmlElement(name)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def generate_document():
    doc = docx.Document()
    
    # Page setup - 1 inch margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Styling helpers
    PRIMARY_COLOR = (0, 51, 102)     # Navy Blue
    SECONDARY_COLOR = (112, 128, 144) # Slate Grey
    TEXT_COLOR = (51, 51, 51)         # Off Black

    def add_heading_1(text):
        h = doc.add_heading('', level=1)
        h.paragraph_format.space_before = Pt(18)
        h.paragraph_format.space_after = Pt(6)
        h.paragraph_format.keep_with_next = True
        run = h.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(18)
        run.bold = True
        run.font.color.rgb = RGBColor(*PRIMARY_COLOR)
        return h

    def add_heading_2(text):
        h = doc.add_heading('', level=2)
        h.paragraph_format.space_before = Pt(14)
        h.paragraph_format.space_after = Pt(4)
        h.paragraph_format.keep_with_next = True
        run = h.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(14)
        run.bold = True
        run.font.color.rgb = RGBColor(*SECONDARY_COLOR)
        return h

    def add_heading_3(text):
        h = doc.add_heading('', level=3)
        h.paragraph_format.space_before = Pt(10)
        h.paragraph_format.space_after = Pt(2)
        h.paragraph_format.keep_with_next = True
        run = h.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(11)
        run.bold = True
        run.font.color.rgb = RGBColor(0, 102, 102)
        return h

    def add_paragraph(text="", bold_prefix="", indent=0):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        if indent > 0:
            p.paragraph_format.left_indent = Inches(indent)
        
        if bold_prefix:
            run_p = p.add_run(bold_prefix)
            run_p.font.name = 'Arial'
            run_p.font.size = Pt(10.5)
            run_p.bold = True
            run_p.font.color.rgb = RGBColor(*TEXT_COLOR)
            
        if text:
            run_t = p.add_run(text)
            run_t.font.name = 'Arial'
            run_t.font.size = Pt(10.5)
            run_t.font.color.rgb = RGBColor(*TEXT_COLOR)
        return p

    def add_bullet(text, bold_prefix=""):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.15
        
        if bold_prefix:
            run_p = p.add_run(bold_prefix)
            run_p.font.name = 'Arial'
            run_p.font.size = Pt(10.5)
            run_p.bold = True
            run_p.font.color.rgb = RGBColor(*TEXT_COLOR)
            
        run_t = p.add_run(text)
        run_t.font.name = 'Arial'
        run_t.font.size = Pt(10.5)
        run_t.font.color.rgb = RGBColor(*TEXT_COLOR)
        return p

    # --- Title Page ---
    p_title_space = doc.add_paragraph()
    p_title_space.paragraph_format.space_before = Pt(72)
    
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("BOSS (BANK OPERATIONS & SCREENING SYSTEM)\n")
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(26)
    run_title.bold = True
    run_title.font.color.rgb = RGBColor(*PRIMARY_COLOR)
    
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run("Comprehensive Codebase Architecture & Reference Document\n")
    run_sub.font.name = 'Arial'
    run_sub.font.size = Pt(14)
    run_sub.bold = False
    run_sub.font.color.rgb = RGBColor(*SECONDARY_COLOR)
    
    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_meta.paragraph_format.space_before = Pt(144)
    run_meta = p_meta.add_run("Client: Karnataka Bank Limited (KBL)\nPrepared for: Development & Audit Teams\nDate: June 2026\nVersion: 1.0 (Detailed Specification)")
    run_meta.font.name = 'Arial'
    run_meta.font.size = Pt(10.5)
    run_meta.font.color.rgb = RGBColor(128, 128, 128)
    
    doc.add_page_break()

    # --- Section 1 ---
    add_heading_1("1. Executive Summary & System Domain")
    add_paragraph("The Bank Operations & Screening System (BOSS) is a web application customized for Karnataka Bank Limited. Its primary objective is to manage the reporting, tracking, and evaluation of Non-Performing Assets (NPAs) and delinquency incidents. BOSS provides structured accountability auditing for loans that have fallen into default, ensuring regulatory compliance and oversight of potential internal lapses.")
    
    add_heading_2("1.1 Core Business Operations")
    add_paragraph("The application serves three distinct operational workflows:")
    add_bullet("The committee assesses the commercial metrics of delinquent accounts, reviews sanctioned limits vs. outstanding balances, verifies collateral values, and initiates early delinquency cases.", bold_prefix="Preliminary Screening Committee (PSC) Review: ")
    add_bullet("Formed to evaluate if the default was caused, enabled, or exacerbated by bank employee negligence, procedural violations, or collusion. This stage evaluates specific staff designations, past branch locations, and lists explicit lapses.", bold_prefix="Staff Accountability Committee (SAC) Review: ")
    add_bullet("The official corporate governance mechanism that compiles decisions, logs attendee lists, tracks review types (e.g., PSC1, PSC2, SAC), and publishes printable meeting files. A meeting must be officially closed before the default statuses are locked in.", bold_prefix="Minutes of Meetings (MOM) Management: ")

    # --- Section 2 ---
    add_heading_1("2. Technical Architecture & Technology Stack")
    add_paragraph("The application is developed as a modular web system utilizing standard Python libraries and Django's model-view-controller abstractions.")
    
    add_heading_2("2.1 Core Frameworks")
    add_bullet("Django 3.1.2 (LTS): Manages ORM persistence, session configurations, form postings, and HTML template rendering.", bold_prefix="Core Framework: ")
    add_bullet("Django REST Framework 3.12.1: Exposes APIs for dynamic UI validations (e.g., branch and user check forms) and dashboard query endpoints.", bold_prefix="API Layer: ")
    add_bullet("django-session-timeout (0.1.0): Configured in settings to automatically expire user sessions after 15 minutes of inactivity (900 seconds) and redirect to the logout route.", bold_prefix="Session Security: ")

    add_heading_2("2.2 Multi-Tenant Database Layout")
    add_paragraph("BOSS supports multi-tenancy to host distinct database catalogs for different bank branches or regions:")
    add_bullet("Mapped through the custom routing module tenants.router.TenantRouter. Individual domains (e.g., fincare.localhost, jana.localhost, hdfc.localhost, kotak.localhost) map to separate SQLite files (default.db.sqlite3).", bold_prefix="Development Environment (dev.py): ")
    add_bullet("Utilizes a single production PostgreSQL instance on port 5945. Search paths are isolated to the specific schema (cbcadm) using connection options settings, protecting transactional data from crossing tenant boundaries.", bold_prefix="Production Environment (prod.py): ")

    add_heading_2("2.3 LDAP Active Directory Authentication")
    add_paragraph("Employee authentication relies on direct Active Directory integration instead of local database password hashes:")
    add_bullet("The ad_login_test function in boss_admin/dbutil.py establishes an LDAP connection over port 389 with the server domain 'ktktest.com'.", bold_prefix="LDAP Bind: ")
    add_bullet("The system binds with the user username format '{domain_id}@ktktest.com' and the supplied password.", bold_prefix="Authentication: ")
    add_bullet("Upon successful authentication, the system searches DC=ktktest,DC=COM under the SUBTREE scope with the filter samaccountname={domain_id} and retrieves: employeeID, givenname, mail, samaccountname, physicalDeliveryOfficeName, title, and Department.", bold_prefix="Attribute Extraction: ")
    add_bullet("The con_current_login function tracks active session keys. If a user logs in from a new IP or device, any older active session matching their emp_id in log_data is terminated, and the previous session key is deleted from session_data.", bold_prefix="Concurrent Logins: ")

    add_heading_2("2.4 Document Management and PDF Output Pipeline")
    add_bullet("Uploaded documents (such as loan approval letters and audit notes) are tracked using DocumentTable. The binary data is stored directly in base64/text format within the database.", bold_prefix="Document Uploads: ")
    add_bullet("The pdf_exports.py module builds dynamic, printable PDFs of PSC reviews, SAC reviews, and MOM logs using ReportLab. It formats report covers, tables, signature sections, page numbers, and embeds PNG/JPG digital signatures from the EmployeeMaster table.", bold_prefix="PDF Generation: ")
    add_bullet("Uses openpyxl and tablib to parse and validate spreadsheet uploads when administrators import bulk employee or branch registries, tracking parsing issues to Excel files.", bold_prefix="Excel Processing: ")

    # --- Section 3 ---
    doc.add_page_break()
    add_heading_1("3. Directory Structure & Key Code Files")
    add_paragraph("Below is an architectural summary mapping key folders and files in the repository to their development roles:")

    headers_dir = ["Directory / File Path", "Role / Component Description", "Key Modules / Files Included"]
    data_dir = [
        ["boss_v1/", "Django configuration root. Manages global URL routers, wsgi hooks, and environment configurations.", "urls.py, wsgi.py, asgi.py"],
        ["boss_v1/settings/", "Environment-specific settings split to isolate local dev databases from production credentials.", "base.py, dev.py, prod.py"],
        ["boss_admin/", "Handles core registries (employees, branches, active directories, permissions, and session audits).", "models.py, dbutil.py, log_utils.py"],
        ["boss_admin/views/", "Controllers managing login/logout, designation matrix configs, and roster actions.", "kbl_auth.py, roles_manage.py, main.py"],
        ["preliminaryscreeningcommittee_review/", "Core operational workflows (case reviews, staff accountability logs, and meeting minutes).", "models.py, urls.py"],
        ["preliminaryscreeningcommittee_review/views/", "Implements Maker/Checker updates, MOM date overlapping checks, and ReportLab PDF layout exports.", "psc_views.py, sac_views.py, pdf_exports.py"],
        ["notifications_app/", "Utility system sending context-driven status update notices to committee members.", "models.py, context_processors.py, views.py"],
        ["CustomLog/ & Logs/", "Log directories separating verbose Django framework logs from specific committee audit logs (psc.log).", "psc.log, registers.log"]
    ]
    
    table_dir = doc.add_table(rows=len(data_dir)+1, cols=3)
    table_dir.style = 'Light Shading Accent 1'
    for idx, name in enumerate(headers_dir):
        cell = table_dir.rows[0].cells[idx]
        cell.text = name
        set_cell_margins(cell)
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.name = 'Arial'
                r.font.size = Pt(9.5)
                
    for r_idx, row_data in enumerate(data_dir):
        for c_idx, val in enumerate(row_data):
            cell = table_dir.rows[r_idx+1].cells[c_idx]
            cell.text = val
            set_cell_margins(cell)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.name = 'Arial'
                    r.font.size = Pt(9.5)

    # --- Section 4 ---
    doc.add_page_break()
    add_heading_1("4. Database Schema & Data Models")
    add_paragraph("The application's models are split into master tables (boss_admin) and transaction-focused tables (screening committee).")
    
    add_heading_2("4.1 Master Models (boss_admin/models.py)")
    
    headers_master = ["Model Name", "DB Table Name", "Description & Key Fields"]
    data_master = [
        ["BranchMaster", "branch_master", "Details all branches. Fields: branch_code (Unique, PK), branch_name, zone, state, retention_limit."],
        ["EmployeeMaster", "employee_master", "Stores employee accounts. Fields: emp_id (Unique PK), domain_id, emp_name, pwd (hash), designation, active, signature (FileField)."],
        ["RegistersMaster", "registers_master", "Catalog of system sections. Fields: registers_code (Unique PK), registers_type, registers_desc, is_active."],
        ["RegistersRoleMaster", "registers_role_master", "Defines roles within specific registers. Fields: role_id (PK), role_name, role_desc, registers_code (FK)."],
        ["EmployeeRegisterRoleMaster", "employee_register_role_master", "Associates employees to specific register roles and branches. Fields: emp_id (FK), registers_code (FK), role_id (FK), branch_code (FK), is_active."],
        ["AuditorConnectionMaster", "auditor_connection_master", "Tracks auditor assignments. Fields: emp_id (FK), registers_code (FK), role_id (FK), branch_code (FK), assigner_branch_code (FK), is_active."],
        ["DesignationMatrix", "designation_matrix", "Maps hierarchical positions to permissions. Fields: role_type, role_name, role_code, designations (JSONField), is_active."]
    ]
    
    table_m = doc.add_table(rows=len(data_master)+1, cols=3)
    table_m.style = 'Light Shading Accent 1'
    for idx, name in enumerate(headers_master):
        cell = table_m.rows[0].cells[idx]
        cell.text = name
        set_cell_margins(cell)
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.name = 'Arial'
                r.font.size = Pt(9.5)
                
    for r_idx, row_data in enumerate(data_master):
        for c_idx, val in enumerate(row_data):
            cell = table_m.rows[r_idx+1].cells[c_idx]
            cell.text = val
            set_cell_margins(cell)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.name = 'Arial'
                    r.font.size = Pt(9.5)

    add_heading_2("4.2 Committee Review Models (preliminaryscreeningcommittee_review/models.py)")
    
    headers_comm = ["Model Name", "DB Table Name", "Description & Key Fields"]
    data_comm = [
        ["MOMTable", "mom_table", "Minutes of Meetings. Fields: id (PK), meeting_id, mom_creation_date, mom_date, audience (JSONField), active, review_type (psc/sac), mom_users_log (JSONField)."],
        ["CustomerTable", "customer_table", "Stores customer records. Fields: cust_id, borrower_name, total_exposure, acc_count, status, npa_status, sac_details (JSONField), lapses_details (JSONField)."],
        ["PSCTable", "psc_table", "Tracks Preliminary Screening Committee reviews. Fields: cust_id, borrower_name, sanc_limit, npa_date, staff_accountability (JSONField), status, current_role, mom_id (FK), psc_users_log (JSONField)."],
        ["CreditFacilityTable", "creditfacility_table", "Lists loan products of default accounts. Fields: psc_id (FK to CustomerTable), credit_feci_slno, lan (Loan Acc Num), balance, outstanding limit, asset_classification, cf_aod_data (JSONField)."],
        ["SecuritiesTable", "securities_table", "Collateral valuation database. Fields: psc_id (FK to CustomerTable), security_nature, lan, latest_valuation, insurance_value, insurance_valid_upto, cersai_num."],
        ["SACTable", "sac_table", "Tracks Staff Accountability reviews. Fields: psc_rec_id (FK to PSCTable), sac_rec_id, emp_name, staff_no, present_designation, lapses_data (JSONField), status, current_role, mom_id (FK), sac_users_log (JSONField)."],
        ["DocumentTable", "document_table", "Saves attachments. Fields: document_id (PK), review_id, review_type, file (TextField for base64 data), file_name, file_type."]
    ]
    
    table_c = doc.add_table(rows=len(data_comm)+1, cols=3)
    table_c.style = 'Light Shading Accent 1'
    for idx, name in enumerate(headers_comm):
        cell = table_c.rows[0].cells[idx]
        cell.text = name
        set_cell_margins(cell)
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.name = 'Arial'
                r.font.size = Pt(9.5)
                
    for r_idx, row_data in enumerate(data_comm):
        for c_idx, val in enumerate(row_data):
            cell = table_c.rows[r_idx+1].cells[c_idx]
            cell.text = val
            set_cell_margins(cell)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.name = 'Arial'
                    r.font.size = Pt(9.5)

    add_heading_2("4.3 Core JSON Schema Definitions")
    add_paragraph("Due to the complex and dynamic nature of loan delinquency reviews, several fields are structured as JSON columns in the database rather than hardcoded tables:")
    
    add_bullet("Stored in CustomerTable. Contains structural information about evaluated bank staff: 'records': [{'staff_no': ..., 'staff_name': ..., 'staff_branch_code': ..., 'staff_designation': ..., 'staff_lapses_details': ...}].", bold_prefix="CustomerTable.lapses_details: ")
    add_bullet("Stored in PSCTable. Holds 20+ distinct validation fields, including: 'pre_sanction', 'post_sanction', 'post_sanction_sa', 'cersai', 'sec_availability', 'cgtmse_loan', 'all_aod' (list of Acknowledgement of Debts), and 'staff_data'.", bold_prefix="PSCTable.staff_accountability: ")
    add_bullet("Tracks audit history of the maker-checker workflow: {'bo_maker': [...], 'bo_checker': [...], 'ro_maker': [...], 'ro_checker': [...], 'ho_maker': [...], 'ho_checker': [...], 'convener': [...]}. Each entry logs 'user', 'ses_key', 'designation', and 'status'.", bold_prefix="PSCTable.psc_users_log / SACTable.sac_users_log: ")

    # --- Section 5 ---
    doc.add_page_break()
    add_heading_1("5. Key Workflows & State Machine Diagrams")
    
    add_heading_2("5.1 Maker/Checker Status Transitions")
    add_paragraph("Every screening case must flow through a structured validation chain. Below is the transition mapping showing the responsible role and status updates:")
    
    headers_state = ["Step Name", "Assignee / Role", "Current Status", "Action Available", "Next Status (Success / Return)"]
    data_state = [
        ["1. Initiate Case", "BO Maker", "Draft", "Submit / Re-submit", "Submitted (to BO Checker)"],
        ["2. Local Audit", "BO Checker", "Submitted", "Approve / Reject", "Approved (to RO Maker) / Rejected (to BO Maker)"],
        ["3. Regional Draft", "RO Maker", "Approved", "Submit Regional Logs", "Submitted (to RO Checker)"],
        ["4. Regional Audit", "RO Checker", "Submitted", "Approve / Reject", "Approved (to HO Maker) / Rejected (to RO or BO Maker)"],
        ["5. HO Appraisal", "HO Maker", "Approved", "Submit HO Logs", "Submitted (to HO Checker)"],
        ["6. HO Audit", "HO Checker", "Submitted", "Approve / Reject", "Approved (to Convener) / Rejected (to HO, RO, or BO Maker)"],
        ["7. Final Closure", "Convener", "Approved", "Associate to MOM / Close", "MOM Closed (Case complete)"]
    ]
    
    table_s = doc.add_table(rows=len(data_state)+1, cols=5)
    table_s.style = 'Light Shading Accent 1'
    for idx, name in enumerate(headers_state):
        cell = table_s.rows[0].cells[idx]
        cell.text = name
        set_cell_margins(cell)
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.name = 'Arial'
                r.font.size = Pt(9.5)
                
    for r_idx, row_data in enumerate(data_state):
        for c_idx, val in enumerate(row_data):
            cell = table_s.rows[r_idx+1].cells[c_idx]
            cell.text = val
            set_cell_margins(cell)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.name = 'Arial'
                    r.font.size = Pt(9.5)

    add_heading_2("5.2 Business Rules for Workflow Rejections")
    add_paragraph("Rejection levels are tracked in a specialized JSON payload (reject_remarks_div) during the POST transaction. When an auditor rejects a review:")
    add_bullet("The rejecting user specifies the target level ('BO' or 'RO' or 'HO').", bold_prefix="Target Resolution: ")
    add_bullet("The system extracts the 'rej_lvl' key and updates current_role back to the corresponding Maker role.", bold_prefix="Role Demotion: ")
    add_bullet("If a Staff Accountability form (SAC) is rejected more than 14 times (tracked by len(reject_remarks) in sac_views.py), the account is flagged as 'Terminated' and removed from the active queue.", bold_prefix="Max Rejections Rule: ")

    add_heading_2("5.3 Minutes of Meetings (MOM) Validation Rules")
    add_paragraph("When a convener schedules an MOM, the system validates the dates in psc_views.py (mom_db_store):")
    add_bullet("The system prevents the creation of active meeting instances if a meeting exists on that exact date.", bold_prefix="Duplicate Date Check: ")
    add_bullet("It ensures a clean schedule by preventing the overlapping of multiple active meetings across past, present, and future ranges.", bold_prefix="Overlap Guard: ")
    add_bullet("When a meeting is finalized, mom_total_count is called to link the evaluated loan records, lock their reviews, and set their status to closed.", bold_prefix="MOM Lock-in: ")

    # --- Section 6 ---
    doc.add_page_break()
    add_heading_1("6. Detailed Walkthrough of Key Operations")
    
    add_heading_2("6.1 Saving PSC Reviews (psc_db_store)")
    add_paragraph("The psc_db_store function in preliminaryscreeningcommittee_review/views/psc_views.py handles case persistence. The execution sequence is detailed below:")
    add_bullet("Accepts the request object and a 'data' array containing basic details, credit facilities, securities, and staff lapses.", bold_prefix="1. Payload Intake: ")
    add_bullet("Examines role_val ('edit1', 'approve2', etc.) and status ('Submitted', 'Approved', 'Rejected') to calculate the next workflow role assignee.", bold_prefix="2. State Routing: ")
    add_bullet("Initializes a django.db.transaction.atomic() context to ensure database consistency during the updates.", bold_prefix="3. Database Transaction: ")
    add_bullet("Queries PSCTable records matching the cust_id, updates general columns (borrower details, exposure, recovery steps), and saves the changes.", bold_prefix="4. PSC Case Update: ")
    add_bullet("Invokes update_PSC_SAC_MOM_logs to append the employee's ID, session token, and timestamp to psc_users_log.", bold_prefix="5. Action Logging: ")
    add_bullet("Updates the matching CustomerTable record's status, assignee, and review ID.", bold_prefix="6. Customer Sync: ")
    add_bullet("Deletes previous creditfacility_table records linked to this customer and loops through the new credit_sanction data to save updated entries.", bold_prefix="7. Credit Facility Refresh: ")
    add_bullet("Deletes previous securities_table records and saves the new collateral valuations.", bold_prefix="8. Securities Refresh: ")
    add_bullet("If documents were uploaded during the session, it extracts their base64 payloads and calls fileUpload.file_upload.insert_documents to save them in DocumentTable.", bold_prefix="9. Document Attachments: ")

    add_heading_2("6.2 Saving SAC Reviews (sac_db_store)")
    add_paragraph("The sac_db_store function operates similarly for Staff Accountability reviews:")
    add_bullet("Pulls basic details including staff_no, present_designation, and present_working.", bold_prefix="1. Payload Intake: ")
    add_bullet("Packs lapses_branch, fraud_element, addr_lett_date, reply_date, and rep_auth_remarks into a lapses_data JSON payload.", bold_prefix="2. Lapse Metadata Packing: ")
    add_bullet("Loops through matching SACTable entries and updates their details, including the lapses_data column.", bold_prefix="3. Database Update: ")
    add_bullet("Appends updates to the sac_users_log column.", bold_prefix="4. Log Sync: ")
    add_bullet("Updates CustomerTable's sac_details JSON field with the new assignee and status, and saves the changes.", bold_prefix="5. Customer Update: ")

    # Save to file
    doc.save("BOSS_Codebase_Guide.docx")
    print("BOSS_Codebase_Guide.docx successfully created.")

if __name__ == "__main__":
    generate_document()
